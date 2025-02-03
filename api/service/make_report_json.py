# langchain 기본 라이브러리

import pandas as pd
import openai
import os
# import pillow_heif
import re
import json
from datetime import datetime

from langchain_openai import ChatOpenAI
from langchain_openai import OpenAIEmbeddings

from langchain_teddynote.messages import stream_response

from langchain_core.prompts import ChatPromptTemplate, FewShotChatMessagePromptTemplate
from langchain_core.example_selectors import (SemanticSimilarityExampleSelector,)
from langchain_community.document_loaders.parsers.pdf import PDFPlumberParser
from langchain_community.document_loaders import PDFMinerLoader

from langchain_chroma import Chroma
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders import UnstructuredPDFLoader
from langchain_community.document_loaders import PDFMinerPDFasHTMLLoader
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain

# word 파일을 만들기 위한 python-docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL


FILE_PATH = "./특이민원보고서_공직자응대매뉴얼.pdf"
file_path = os.path.join(os.getcwd(), "api_key.txt")

# 1. API 키 읽기 및 설정
def load_api_key(file_path):
    with open(file_path, 'r') as file:
        return file.read().strip()
    
def selecting_model(api_key):
    os.environ["OPENAI_API_KEY"] = api_key
    llm = ChatOpenAI(temperature=0, model_name="gpt-4o")
    return llm

api_key = load_api_key("api_key.txt")
llm = selecting_model(api_key)
chroma = Chroma("fewshot_chat", OpenAIEmbeddings())

# Load the PDF file
class LoadPdfFile:
    def __init__(self, file_path):
        self.file_path = file_path
        self.data = self.load_pdf_file()
        self.llm = selecting_model(api_key)
        
    def load_pdf_file(self):
        loader = PDFMinerLoader(self.file_path)
        data = loader.load()
        return data

    # 표 데이터 감지 및 추출
    def extract_tables_from_text(self, text):
        tables = []
        rows = text.split("\n")
        for row in rows:
            # 간단한 정규식을 사용해 테이블 형식 감지 (예: 열 구분자가 탭이나 공백인 경우)
            if re.match(r"(\w+\s+)+", row):
                tables.append(row.split())
        return tables
    
    def extract_tables_data(self):
        all_tables = []
        for doc in self.data:
            text_content = doc.page_content  # 페이지 내용
            tables = self.extract_tables_from_text(text_content)
            all_tables.extend(tables)
        return all_tables
    
    def make_prompt(self):
        table_extraction_prompt = """
        Extract all tables from the following text and format them as JSON:
        {text}
        """
        return PromptTemplate(input_variables=["text"], template=table_extraction_prompt)
        
    def make_llm_json(self):
        # LLM 초기화 및 체인 생성
        llm_chain = LLMChain(llm=self.llm, prompt=self.make_prompt())
        for doc in self.data:
            result_text = llm_chain.run({"text": doc.page_content})
            
            try:
                json_match = re.search(r"```json(.*?)```", result_text, re.DOTALL)
                if json_match:
                    json_text = json_match.group(1).strip()
                    result = json.loads(json_text)  # JSON 변환
                    report = result.get("특이민원 발생보고서", {})
                    # JSON 데이터에서 "특이민원 발생보고서" 추출
                    return report
                else: print("JSON 형식을 감지하지 못했습니다. LLM 출력:", result_text)
            except json.JSONDecodeError as e:
                print("JSON 변환 중 오류 발생:", e)
                print("LLM 출력:", result_text)
            except KeyError as e:
                print("JSON에서 키를 찾을 수 없음:", e)

class MakeReport:
    def __init__(self, file_path):
        self.doc = Document()
        loader = LoadPdfFile(file_path)
        self.report = loader.make_llm_json()

    # 폰트 설정 함수
    def set_korean_font(self, paragraph, font_name="맑은 고딕", font_size=11):
        for run in paragraph.runs:
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:eastAsia"), font_name)
            rPr.append(rFonts)
            run.font.size = Pt(font_size)
    
    def add_title(self, text):
        title = self.doc.add_paragraph()
        run = title.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
    def add_table(self):
    # 표 생성 (행: 8, 열: 9)
        table = self.doc.add_table(rows=8, cols=9)
        table.style = 'Table Grid'

        # 첫 번째 행: 발생일자, 부서, 부서장
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "발생일자"
        hdr_cells[1].text = " "
        hdr_cells[2].text = "부서"
        hdr_cells[3].text = " "
        hdr_cells[4].text = "부서장"
        hdr_cells[5].text = " "
        hdr_cells[6].merge(hdr_cells[8])  # 빈 공간 병합

        # 두 번째 행: 특이민원 상황
        row_1_cells = table.rows[1].cells
        row_1_cells[0].text = "특이민원 상황"
        row_1_cells[1].text = "폭언"
        row_1_cells[2].text = "협박"
        row_1_cells[3].text = "폭행"
        row_1_cells[4].text = "성희롱"
        row_1_cells[5].text = "기물파손"
        row_1_cells[6].text = "위험물 소지"
        row_1_cells[7].text = "주취 소란"
        row_1_cells[8].text = "기타"

        # 세 번째 행: 특이민원 체크박스
        row_2_cells = table.rows[2].cells
        for i in range(9):
            row_2_cells[i].text = " "

        # 네 번째 행: 민원인
        row_3_cells = table.rows[3].cells
        row_3_cells[0].text = "민원인"
        row_3_cells[1].text = " "
        row_3_cells[2].text = "전화번호"
        row_3_cells[3].merge(row_3_cells[8])  # 나머지 공간 병합

        # 다섯 번째 행: 담당자 정보
        row_4_cells = table.rows[4].cells
        row_4_cells[0].text = "담당자"
        row_4_cells[1].text = "전화번호"
        row_4_cells[2].text = "담당업무"
        row_4_cells[3].merge(row_4_cells[8])

        # 여섯 번째 행: 특이민원 발생요지
        row_5_cells = table.rows[5].cells
        row_5_cells[0].merge(row_5_cells[8])
        row_5_cells[0].text = "특이민원 발생요지"

        # 일곱 번째 행: 담당자 의견
        row_6_cells = table.rows[6].cells
        row_6_cells[0].merge(row_6_cells[8])
        row_6_cells[0].text = "담당자 의견"

        # 여덟 번째 행: 부서장 의견
        row_7_cells = table.rows[7].cells
        row_7_cells[0].merge(row_7_cells[8])
        row_7_cells[0].text = "부서장 의견"

        # 빈 공간과 정렬 설정
        for row in table.rows:
            for cell in row.cells:
                if not cell.text.strip():  # 빈 셀은 공백으로 설정
                    cell.text = " "
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
    def make_report_detail(self):
        self.add_title("특이민원 발생보고서")
        self.add_table()
    
    def report_save(self):
        time = datetime.now().strftime("%y%m%d_%H%M")
        output_file = f"특이민원_보고서_{time}.docx"
        self.doc.save(output_file)
        print(f"문서가 {output_file}에 저장되었습니다.")