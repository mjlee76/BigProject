from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from langchain.schema import HumanMessage

from pydantic import BaseModel

import basic_module as bm
import datetime
import os
import requests


#db에서 받아올 정보
department = "민원복지과"
department_supervisior = "홍길동"
label = '성희롱'
customer = "김태환"
customer_Telephone = "010-1234-4567"
manager = "김영희"
manager_telephone = "02-873-4466"
manager_task = "복지민원 담당"

label_mapping = {
                0: "정상", 1: "악성", 2: "욕설",
                3: "외모", 4: "장애인", 5: "인종",
                6: "종교", 7: "지역", 8: "성차별",
                9: "나이", 10: "협박", 11: "성희롱"}

class UserInfo(BaseModel):
    user_id : str
    user_name: str
    gender : str
    role : str
    birth_date : str
    telephone : str
    address : str
    email: str
    create_date : str
    count : int
    
class PostBody(BaseModel):
    title: str
    content: str
    user: UserInfo
    
class ReportBody(BaseModel):
    category : str
    report_path : str
    create_date : object
    
path = "./api_key.txt"
api_key = bm.load_api_key(path)
os.environ["OPENAI_API_KEY"] = api_key
llm = bm.selecting_model(api_key)

class MakeReport():
    def __init__(self):
        self.doc = Document("C:/Users/User/Desktop/빅프로젝트/BigProject/api/service/특이민원_발생보고서.docx")
    
    def set_font(paragraph, font_name="맑은 고딕", font_size=12):
        for run in paragraph.runs:
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:eastAsia"), font_name)
            rPr.append(rFonts)
            run.font.size = Pt(font_size)

    def report_prompt(self, post_body : PostBody):
        title = post_body.title
        content = post_body.content
        full_text = f"제목 : {title}" + " " + f"내용 : {content}"
        prompt = (
            "당신은 민원에 대한 처리를 하는 상담사입니다. 특이민원이 발생하여 이에 대한 보고서를 작성해야합니다."
            "다음 문장을 보고 특이민원 발생요지에 대해서 6하원칙에 따라 핵심내용 위주만 간략하게 작성해주세요"
            f"판단할 내용 : {full_text}"
            )
        try:
            # LLM을 사용하여 분류 수행
            response = llm([HumanMessage(content=prompt)])
            return response.content.strip()  # LLM 응답 반환
        except Exception as e:
            print(f"Error: {str(e)}")
            return "오류 발생"

    # 표 검색 후 특정 셀 찾기
    def cell_fill(self, post_body : PostBody, report_body : ReportBody):
        table = self.doc.tables[0]
        table.cell(0, 1).text = "2025-01-01"  # 예: 발생일자
        table.cell(0, 4).text = f"{department}"
        table.cell(0, 7).text =  f"{department_supervisior}"
        table.cell(4, 1).text = f"{customer}"
        table.cell(4, 4).text = f"{customer_Telephone}"
        table.cell(5, 1).text = f"{manager}"
        table.cell(5, 4).text = f"{manager_telephone}"
        table.cell(5, 7).text = f"{manager_task}"
        
        label = report_body.category
        if '성희롱' in label:
            cell = table.cell(2, 4)
        elif '협박' in label:
            cell = table.cell(2, 2)
        elif '악성' in label: 
            cell = table.cell(2, 1)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run("○")
        run.font.size = Pt(22)
        
        prompt = self.report_prompt(post_body)
        table.cell(6, 1).text = prompt
        
    def report_save(self):
        time = datetime.datetime.now().strftime("%y%m%d_%H%M")
        output_file = f"특이민원_보고서_{time}.docx"
        report_file_path = "C:/Users/User/Desktop/빅프로젝트/BigProject/tellMe/tellMe-reports/"
        self.doc.save(report_file_path + output_file)
        print(f"문서가 {output_file}에 저장되었습니다.")
        return time, output_file
