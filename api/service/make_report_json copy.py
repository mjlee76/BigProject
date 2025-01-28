#!pip install langchain_chroma
#!pip install python-docx
#!pip install openai
#!pip install langchain_openai
#!pip install langchain
#!pip install langchain_teddynote
#!pip install -qU pypdf
#!pip install langchain_community
#!pip install -qU unstructured
#!pip install pdfminer
#!pip install -qU pymupdf
#!pip install -qU rapidocr-onnxruntime
#!pip install -qU pypdf
#!pip install pdfminer.six==20221105
#!pip install pillow-heif
#!pip install pdfplumber langchain

# langchain 기본 라이브러리

import pandas as pd
import openai
import os
import pillow_heif
import re
import json

from langchain_openai import ChatOpenAI
from langchain_openai import OpenAIEmbeddings

from langchain_teddynote.messages import stream_response

from langchain_core.prompts import ChatPromptTemplate, FewShotChatMessagePromptTemplate
from langchain_core.example_selectors import (SemanticSimilarityExampleSelector,)
from langchain_community.document_loaders.parsers.pdf import PDFPlumberParser
from langchain_community.document_loaders import PDFMinerLoader

from langchain_chroma import Chroma
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.document_loaders import UnstructuredPDFLoader
from langchain_community.document_loaders import PDFMinerPDFasHTMLLoader
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain

# word 파일을 만들기 위한 python-docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 1. API 키 읽기
def load_api_key(file_path):
    with open(file_path, 'r') as file:
        return file.read().strip()

# 2. API 키 설정
api_key = load_api_key("api_key.txt")
os.environ["OPENAI_API_KEY"] = api_key

llm = ChatOpenAI(temperature=0, model_name="gpt-4o",)

#벡터 DB 연결
chroma = Chroma("fewshot_chat", OpenAIEmbeddings())

# Load the PDF file
FILE_PATH = "./특이민원보고서_공직자응대매뉴얼.pdf"

'''parser = PDFPlumberParser()
processed_data = parser.parse_file(FILE_PATH)
print(processed_data)'''

# PDFMinerPDFasHTMLLoader 인스턴스 생성
loader = PDFMinerPDFasHTMLLoader(FILE_PATH)
# 문서 로드
docs = loader.load()
# 문서의 내용 출력
#print(docs[0])

# 문서 만들어주기

# 문서 생성
doc = Document()

# 제목 추가
title = doc.add_heading("특이민원 발생보고서", level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 폰트 설정 함수
def set_korean_font(paragraph, font_name="맑은 고딕", font_size=11):
    for run in paragraph.runs:
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), font_name)
        rPr.append(rFonts)
        run.font.size = Pt(font_size)

# 테이블 생성
table = doc.add_table(rows=0, cols=3, style="Table Grid")

# 첫 번째 행 추가 (헤더)
header_row = table.add_row().cells
header_row[0].text = "항목"
header_row[1].text = "세부 항목"
header_row[2].text = "내용"

# 항목 추가
for key, value in report.items():
    if isinstance(value, dict):  # 특이민원유형 및 담당자 정보
        for sub_key, sub_value in value.items():
            row = table.add_row().cells
            row[0].text = key
            row[1].text = sub_key
            row[2].text = str(sub_value)
    else:  # 일반 항목
        row = table.add_row().cells
        row[0].text = key
        row[1].text = ""
        row[2].text = str(value)

# 문서 저장
output_file = "특이민원_보고서_v2.docx"
doc.save(output_file)

print(f"문서가 {output_file}에 저장되었습니다.")