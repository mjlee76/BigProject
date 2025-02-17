import os
import re
import torch
import asyncio
import aiofiles
import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from langchain.schema import HumanMessage

#Loader 불러오기
from langchain_community.document_loaders.parsers.pdf import PDFPlumberParser
from langchain_teddynote.document_loaders import HWPLoader
from langchain_community.document_loaders import (
PyPDFLoader, Docx2txtLoader, TextLoader, DirectoryLoader, UnstructuredPDFLoader, 
UnstructuredPDFLoader, PDFMinerPDFasHTMLLoader, PDFMinerLoader)

from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage
from langchain_chroma import Chroma
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain

from pydantic import BaseModel
from transformers import AutoTokenizer, AutoModelForSequenceClassification

#1. api 키 불러오기 및 설정정
async def load_api_key(file_path):
    async with aiofiles.open(file_path, 'r') as file:
        return (await file.read()).strip()

#2. 모델 설정
async def selecting_model(api_key):
    os.environ["OPENAI_API_KEY"] = api_key
    llm = await asyncio.to_thread(ChatOpenAI, temperature=0, model_name="gpt-4")
    return llm

#3. 모델 프롬프트 출력
class TextClassifier:
    def __init__(self):
        model_path = "./20250204_roberta 파인튜닝_final"
        self.model = AutoModelForSequenceClassification.from_pretrained(model_path)
        self.tokenizer = AutoTokenizer.from_pretrained("klue/roberta-base")

    def split_text(self, text, max_length):
        """텍스트를 토큰 길이를 기준으로 나눔"""
        tokens = self.tokenizer(text, truncation=False, padding=False)
        input_ids = tokens["input_ids"]
        chunks = [input_ids[i:i + max_length] for i in range(0, len(input_ids), max_length)]
        # 토큰 ID를 다시 텍스트로 디코딩
        return [self.tokenizer.decode(chunk, skip_special_tokens=True) for chunk in chunks if len(chunk) > 0]

    def classify_text(self, text):
        max_length = 512
        chunks = self.split_text(text, max_length - 2)
        results = [] 
        for chunk in chunks:
            try:
                inputs = self.tokenizer(chunk, truncation=True, padding="max_length", max_length=max_length, return_tensors="pt")
                with torch.no_grad():
                    outputs = self.model(**inputs)
                    logits = outputs.logits
                    probabilities = torch.sigmoid(logits).squeeze()  # Sigmoid로 확률 계산
                    results.append(probabilities)
            except Exception as e:
                print(f"Error processing chunk: {chunk}")
                print(f"Error details: {e}")

        if results:
            final_probabilities = torch.mean(torch.stack(results), dim=0)  # 평균 확률
            threshold = 0.5
            predicted_labels = (final_probabilities > threshold).nonzero(as_tuple=True)[0].tolist()

            # 라벨 매핑
            label_mapping = {
                0: "정상", 1: "모욕", 2: "욕설",
                3: "외모차별", 4: "장애인차별", 5: "인종차별",
                6: "종교차별", 7: "지역차별", 8: "성차별",
                9: "나이차별", 10: "협박", 11: "성희롱",
            }

            # 결과 출력
            predicted_labels = [label_mapping[label] for label in predicted_labels]
            predicted_labels_text = ','.join(f'{label}' for label in predicted_labels)
            return predicted_labels_text
        else:
            print("텍스트 조각 처리 중 문제가 발생하여 결과를 생성할 수 없습니다.")
            predicted_labels_text = ''
            return predicted_labels_text

#4. 분류
def extract_label(text):
    if isinstance(text, list) and len(text) == 1:
        return text[0]
    return text

#5. 순화
class ChangeText:
    
    def __init__(self):
        self.llm = None  # 나중에 async 초기화에서 설정할 llm

    async def init(self):
        file_path = os.path.join(os.getcwd(), "api_key.txt")
        api_key = await load_api_key(file_path)
        self.llm = await selecting_model(api_key)
    
    def change_generate_prompt(self, text):
        return (
            f"다음 텍스트를 읽고, 정상이 아닌 텍스트를 정중하고 부드러운 표현으로 순화해주세요:\n"
            f"{text}\n\n"
        )
    
    async def change_text(self, text):
        prompt = self.change_generate_prompt(text)
        try:
            # LLM을 사용하여 분류 수행
            response = await asyncio.to_thread(lambda: self.llm([HumanMessage(content=prompt)]))
            return response.content.strip()  # LLM 응답 반환
        except Exception as e:
            return f"Error: {str(e)}"

#6. 법령 경고문 출력
def laws_caution(label):
    print()

class UserInfo(BaseModel):
    user_name: str
    phone : str
    count : int

# 게시글 정보
class PostBody(BaseModel):
    question_id : int
    title: str
    content: str
    user: UserInfo

#DB로 넘길 report 정보
class ReportBody(BaseModel):
    category : list
    post_origin_data : dict
    file_name : str = ""
    create_date : str
    
class CombinedModel(BaseModel):
    post_data: PostBody
    report_req: ReportBody

department = "민원복지과"
department_supervisior = "홍길동"
label = '성희롱'
manager = "김영희"
manager_telephone = "02-873-4466"
manager_task = "복지민원 담당"

class MakeReport():
    def __init__(self):
        self.doc = Document(os.path.join(os.getcwd(), "특이민원_발생보고서.docx"))
        self.llm = None
        self.time = datetime.datetime.now().strftime("%y%m%d_%H%M%S")

    async def init(self):
        # 비동기 초기화 메서드: API 키 로드와 모델 초기화
        self.path = os.getcwd()
        file_path = os.path.join(os.getcwd(), "api_key.txt")
        api_key = await load_api_key(file_path)
        self.llm = await selecting_model(api_key)
    
    def set_font(paragraph, font_name="맑은 고딕", font_size=12):
        for run in paragraph.runs:
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:eastAsia"), font_name)
            rPr.append(rFonts)
            run.font.size = Pt(font_size)

    def report_prompt(self, report_body : ReportBody):
        title = report_body.post_origin_data["제목"]
        content = report_body.post_origin_data["내용"]
        full_text = f"제목 : {title}" + " " + f"내용 : {content}"
        prompt = (
            "당신은 민원에 대한 처리를 하는 상담사입니다. 특이민원이 발생하여 이에 대한 보고서를 작성해야합니다."
            "다음 문장을 보고 특이민원 발생요지에 대해서 6하원칙에 따라 핵심내용 위주만 간략하게 작성해주세요"
            "예) 폭언, 욕설이 담겨져있는 내용 민원글 작성"
            f"판단할 내용 : {full_text}"
            )
        try:
            # LLM을 사용하여 분류 수행
            response = self.llm([HumanMessage(content=prompt)])
            return response.content.strip()  # LLM 응답 반환
        except Exception as e:
            print(f"Error: {str(e)}")
            return "오류 발생"

    # 표 검색 후 특정 셀 찾기
    def cell_fill(self, post_body : PostBody, report_body : ReportBody):
        table = self.doc.tables[0]
        table.cell(0, 1).text = f"{self.time}"  # 예: 발생일자
        table.cell(4, 1).text = f"{post_body.user.user_name}"
        table.cell(4, 4).text = f"{post_body.user.phone}"
        table.cell(0, 4).text = f"{department}"
        table.cell(0, 7).text = f"{department_supervisior}"
        table.cell(5, 1).text = f"{manager}"
        table.cell(5, 4).text = f"{manager_telephone}"
        table.cell(5, 7).text = f"{manager_task}"
        
        label = report_body.category
        print(label)
        if '성희롱' in label:
            cell_label = table.cell(2, 4)
        elif '협박' in label:
            cell_label = table.cell(2, 2)
        elif '폭언' in label or '욕설' in label: 
            cell_label = table.cell(2, 1)
        else:
            cell_label = table.cell(2, 8)
        paragraph = cell_label.paragraphs[0]
        run = paragraph.add_run("○")
        run.font.size = Pt(22)
        
        prompt = self.report_prompt(report_body)
        table.cell(6, 1).text = prompt
        
    def report_save(self, post_body : PostBody):
        output_file = f"{post_body.question_id}번_게시글_특이민원_보고서.docx"
        report_file_path = "C:/Users/User/BigProject/tellMe/tellMe-reports"
        self.doc.save(f"{report_file_path}" + f"/{output_file}")
        print(f"문서가 {output_file}에 저장되었습니다.")
        return self.time, output_file

class LoadDocumentFile:
    def __init__(self):
        self.llm = None
        self.llm_cla = None

    async def init(self):
        file_path = os.path.join(os.getcwd(), "api_key.txt")
        api_key = await load_api_key(file_path)
        self.llm = await self.selecting_model(api_key)
        self.llm_cla = await self.selecting_classify_model(api_key)
    
    async def selecting_model(self, api_key):
        os.environ["OPENAI_API_KEY"] = api_key
        llm = await asyncio.to_thread(ChatOpenAI, temperature=0, model_name="gpt-4")
        return llm
    
    async def selecting_classify_model(self, api_key):
        os.environ["OPENAI_API_KEY"] = api_key
        llm = await asyncio.to_thread(ChatOpenAI, temperature=0, model_name="gpt-4o")
        return llm
    
    #loader로 data 저장하기
    async def select_loader(self, docu_file_path):
        if os.path.isdir(docu_file_path):
            filenames = os.listdir(docu_file_path)
            if not filenames:
                raise ValueError(f"디렉터리 '{docu_file_path}'에 파일이 없습니다.")
        # 예시로 첫 번째 파일을 사용
            docu_path = os.path.join(docu_file_path, filenames[0])
        else:
        # 단일 파일 경로인 경우 그대로 사용
            docu_path = docu_file_path
        ext = os.path.splitext(docu_path)[1].lower()
        # 조건문을 사용하여 확장자에 따라 loader 선택
        if ext == ".pdf":
            loader = PyPDFLoader(docu_path)
        elif ext in [".hwp", ".hwpx"]:
            loader = HWPLoader(docu_path)
        elif ext in [".doc", ".docx"]:
            loader = Docx2txtLoader(docu_path)
        elif ext == ".txt":
            loader = TextLoader(docu_path, autodetect_encoding=True)
        else: 
            raise ValueError(f"지원되지 않는 파일 형식: {ext}")
        data = loader.load()
        return data
    
    async def make_prompt(self):
        prompt = """
        주어진 텍스트에서 주요 항목을 추출한 뒤,
        JSON 형태로 만들어 주세요.
        
        텍스트:
        {text}
        """
        return PromptTemplate(input_variables=["text"], template=prompt)
        
    async def make_llm_text(self, data):
        # LLM 초기화 및 체인 생성
        prompt = await self.make_prompt()
        llm_chain = LLMChain(llm=self.llm, prompt=prompt)
        tasks = []
        for doc in data:
            task = asyncio.to_thread(llm_chain.run, {"text": doc.page_content})
            tasks.append(task)
        results = await asyncio.gather(*tasks)
        # 결과 후처리: 개행, 공백 정리 등
        cleaned_results = [re.sub(r"\s+", " ", r.replace("\n", " ")).strip() for r in results]
        return cleaned_results
    
    async def make_classify_prompt(self):
        return PromptTemplate(
        input_variables=["cleaned_results", "label_mapping"],
        template="""
        당신은 악성민원을 탐지하는 AI 어시스턴트입니다.
        다음 텍스트를 확인하고, 만약 **실제로 악의적/모욕적/차별적**이라면
        아래 제공된 라벨 중 어느 것에 해당하는지 판단해주세요.

        - 단순히 “폭언”, “협박”, “성희롱” 같은 단어를 언급하는 **양식/메뉴얼** 설명이거나, 
        문서가 그런 표현을 **사례나 주의사항**으로 나열하는 것만으로는 악성으로 보지 않습니다.
        - 텍스트가 **직접 공격적**이거나 **누군가를 모욕**, **차별**, **협박**, **성희롱**하고 있다면 해당 라벨을 골라주세요.
        - 그 외에는 "정상"이라고 답변해주세요.
        텍스트: {cleaned_results}
        라벨 목록: {label_mapping}
        
        아래 형식으로 답변을 해주세요:
        ※ 예) 결과 : 성희롱
        """
        )
    
    async def make_classify_text(self, cleaned_results):
        # LLM 초기화 및 체인 생성
        label_mapping = ["정상", "모욕","욕설", "외모차별", "장애인차별", "인종차별",
                "종교차별", "지역차별","성차별", "나이차별", "협박", "성희롱"]
        prompt = await self.make_classify_prompt()
        llm_chain = LLMChain(llm=self.llm_cla, prompt=prompt)
        result = asyncio.to_thread(llm_chain.run, {"cleaned_results": cleaned_results, "label_mapping": ", ".join(label_mapping)})
        result_text = await result
        label = str(result_text.split(':', 1)[1].strip())
        return label
        

        
    
    