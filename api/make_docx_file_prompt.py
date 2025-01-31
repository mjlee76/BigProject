from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

# 문서 생성
doc = Document()

# 제목 추가
def add_title(doc, text):
    title = doc.add_paragraph()
    run = title.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_table(doc):
    # 표 생성 (행: 8, 열: 9)
    table = doc.add_table(rows=8, cols=9)
    table.style = 'Table Grid'

    # 첫 번째 행: 발생일자, 부서, 부서장
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "발생일자"
    hdr_cells[1].text = " "
    hdr_cells[2].text = "부서"
    hdr_cells[3].text = " "
    hdr_cells[4].text = "부서장"
    hdr_cells[5].text = " "
    hdr_cells[6].merge(hdr_cells[8])  # 빈 공간 병합

    # 두 번째 행: 특이민원 상황
    row_1_cells = table.rows[1].cells
    row_1_cells[0].text = "특이민원 상황"
    row_1_cells[1].text = "폭언"
    row_1_cells[2].text = "협박"
    row_1_cells[3].text = "폭행"
    row_1_cells[4].text = "성희롱"
    row_1_cells[5].text = "기물파손"
    row_1_cells[6].text = "위험물 소지"
    row_1_cells[7].text = "주취 소란"
    row_1_cells[8].text = "기타"

    # 세 번째 행: 특이민원 체크박스
    row_2_cells = table.rows[2].cells
    for i in range(9):
        row_2_cells[i].text = " "

    # 네 번째 행: 민원인
    row_3_cells = table.rows[3].cells
    row_3_cells[0].text = "민원인"
    row_3_cells[1].text = " "
    row_3_cells[2].text = "전화번호"
    row_3_cells[3].merge(row_3_cells[8])  # 나머지 공간 병합

    # 다섯 번째 행: 담당자 정보
    row_4_cells = table.rows[4].cells
    row_4_cells[0].text = "담당자"
    row_4_cells[1].text = "전화번호"
    row_4_cells[2].text = "담당업무"
    row_4_cells[3].merge(row_4_cells[8])

    # 여섯 번째 행: 특이민원 발생요지
    row_5_cells = table.rows[5].cells
    row_5_cells[0].merge(row_5_cells[8])
    row_5_cells[0].text = "특이민원 발생요지"

    # 일곱 번째 행: 담당자 의견
    row_6_cells = table.rows[6].cells
    row_6_cells[0].merge(row_6_cells[8])
    row_6_cells[0].text = "담당자 의견"

    # 여덟 번째 행: 부서장 의견
    row_7_cells = table.rows[7].cells
    row_7_cells[0].merge(row_7_cells[8])
    row_7_cells[0].text = "부서장 의견"

    # 빈 공간과 정렬 설정
    for row in table.rows:
        for cell in row.cells:
            if not cell.text.strip():  # 빈 셀은 공백으로 설정
                cell.text = " "
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# 문서 작성
add_title(doc, "특이민원 발생보고서")
add_table(doc)

# 파일 저장
doc.save("특이민원_발생보고서.docx")